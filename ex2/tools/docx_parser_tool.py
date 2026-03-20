from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from typing import Dict, Any, List
import os

COLOR_MAP = {
    '0000FF': 'blue', '00B050': 'green', '008000': 'green', 'FF0000': 'red'
}
HIGHLIGHT_MAP = {
    WD_COLOR_INDEX.YELLOW: 'yellow',
    WD_COLOR_INDEX.BRIGHT_GREEN: 'green',
}

def _font_color_name(run) -> str | None:
    color = None
    try:
        rgb = run.font.color.rgb
        if rgb is not None:
            color = COLOR_MAP.get(str(rgb).upper(), str(rgb).upper())
    except Exception:
        pass
    return color


def _highlight_name(run) -> str | None:
    try:
        val = run.font.highlight_color
        return HIGHLIGHT_MAP.get(val, str(val) if val is not None else None)
    except Exception:
        return None


def parse_docx(path: str) -> Dict[str, Any]:
    doc = Document(path)
    paragraphs: List[Dict[str, Any]] = []
    for i, p in enumerate(doc.paragraphs):
        runs = []
        for j, r in enumerate(p.runs):
            runs.append({
                'run_index': j,
                'text': r.text,
                'font_color': _font_color_name(r),
                'highlight_color': _highlight_name(r),
                'bold': bool(r.bold),
                'italic': bool(r.italic),
                'underline': bool(r.underline),
            })
        paragraphs.append({
            'paragraph_index': i,
            'style': p.style.name if p.style else None,
            'text': p.text,
            'runs': runs,
        })
    tables = []
    for ti, t in enumerate(doc.tables):
        rows = []
        for ri, row in enumerate(t.rows):
            row_cells = []
            for ci, cell in enumerate(row.cells):
                row_cells.append({'row': ri, 'col': ci, 'text': '\n'.join(p.text for p in cell.paragraphs).strip()})
            rows.append(row_cells)
        tables.append({'table_index': ti, 'rows': rows})
    return {
        'source_file': os.path.basename(path),
        'paragraph_count': len(paragraphs),
        'table_count': len(tables),
        'paragraphs': paragraphs,
        'tables': tables,
    }

    