import json
from pathlib import Path
from docx import Document

def build_final_sap_docx(
    template_docx_path: str, 
    update_units_path: str,
    draft_json_path: str, 
    output_docx_path: str
) -> None:
    """
    Reads the template, original update units, and AI drafts,
    replaces placeholders, and saves the final DOCX file.
    """
    print(f"Loading template from: {template_docx_path}")
    doc = Document(template_docx_path)
    
    try:
        with open(update_units_path, 'r', encoding='utf-8') as f:
            update_units = json.load(f)
        with open(draft_json_path, 'r', encoding='utf-8') as f:
            draft_updates = json.load(f)
    except Exception as e:
        print(f"Error loading JSON data: {e}")
        return

    # 1. Create a quick lookup dictionary: { "U0001": "Drafted text..." }
    draft_lookup = {item["unit_id"]: item.get("proposed_sap_text", "") for item in draft_updates}

    replacement_count = 0

    # 2. Iterate through the original update units to do the replacement
    for unit in update_units:
        unit_id = unit.get("unit_id")
        original_text = unit.get("template_text", "") # FIXED KEY
        p_index = unit.get("paragraph_index") # GET PARAGRAPH INDEX
        run_indices = unit.get("run_indices") # Handles the new merged runs list
        
        if unit_id in draft_lookup and original_text:
            drafted_text = draft_lookup[unit_id]
            
            # Replace in the specific normal paragraph if index is available
            if p_index is not None and p_index < len(doc.paragraphs):
                para = doc.paragraphs[p_index]
                
                # If we have specific runs (from blue/green text), preserve formatting
                if run_indices:
                    first_run_idx = run_indices[0]
                    # 1. Put the new text in the FIRST run
                    if first_run_idx < len(para.runs):
                        para.runs[first_run_idx].text = drafted_text
                    
                    # 2. Clear out the rest of the merged runs so they don't duplicate
                    for r_idx in run_indices[1:]:
                        if r_idx < len(para.runs):
                            para.runs[r_idx].text = ""
                            
                    replacement_count += 1
                
                # Fallback for bracket placeholders (no specific runs known)
                else:
                    if original_text in para.text:
                        # ADD ", 1" to only replace the first occurrence (fixes ghosting overlap)
                        para.text = para.text.replace(original_text, drafted_text, 1)
                        replacement_count += 1
            else:
                # Fallback to global search if it's in a table or index is missing
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if original_text in para.text:
                                    # ADD ", 1" to only replace the first occurrence
                                    para.text = para.text.replace(original_text, drafted_text, 1)
                                    replacement_count += 1

    # Save the final document
    output_path = Path(output_docx_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    
    print(f"Success! Replaced {replacement_count} placeholders.")
    print(f"Final SAP draft saved to: {output_path}")

if __name__ == "__main__":
    base_dir = Path(__file__).resolve().parent
    data_dir = base_dir / "data"
    output_dir = base_dir / "outputs"
    
    template_path = str(data_dir / "sap_template.docx")
    update_units_path = str(output_dir / "update_units.json")
    draft_json_path = str(output_dir / "drafted_sap_updates.json")
    final_docx_path = str(output_dir / "Draft_SAP_v1.docx")
    
    if Path(draft_json_path).exists() and Path(update_units_path).exists():
        build_final_sap_docx(template_path, update_units_path, draft_json_path, final_docx_path)
    else:
        print("Missing JSON files. Run the AI workflow first.")